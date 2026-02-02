#!/usr/bin/env python3
"""
Oja POS — Setup & Reference Guide (.docx generator)

Run:  python3 generate-oja-docs.py
Out:  /Users/shile/Documents/Oja POS - Setup & Reference Guide.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.expanduser("~/Documents/Oja POS - Setup & Reference Guide.docx")
BRAND = RGBColor(0xE0, 0x5E, 0x1B)  # Burnt Orange
GRAY = RGBColor(0x66, 0x66, 0x66)
LAST_UPDATED = "February 01, 2026"
VERSION = "1.0"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(sh)


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header row shading."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, 'F5E6D8')
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for ri_idx in range(len(t.rows)):
            for ci_idx, w in enumerate(col_widths):
                t.rows[ri_idx].cells[ci_idx].width = Inches(w)
    doc.add_paragraph()  # spacer
    return t


def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = BRAND
        r.font.size = Pt(14)


def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = BRAND
        r.font.size = Pt(13)


def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


def numbered(doc, text):
    doc.add_paragraph(text, style='List Number')


def normal(doc, text):
    doc.add_paragraph(text)


def code_block(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)


def bold_normal(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


# ── document ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # -- Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Title page ───────────────────────────────────────────────────────
    doc.add_paragraph()  # spacer
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('OJA POS')
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Setup, Configuration & Troubleshooting Guide')
    r.font.size = Pt(16)
    r.font.color.rgb = GRAY

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Last Updated: {LAST_UPDATED}')
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Version {VERSION}')
    r.font.size = Pt(12)

    # ── Table of Contents ────────────────────────────────────────────────
    doc.add_paragraph()
    heading1(doc, 'Table of Contents')
    toc_items = [
        '1. Project Overview',
        '2. Technology Stack',
        '3. Repository & Codebase',
        '4. Domain & DNS Setup',
        '5. Vercel Deployment',
        '6. Supabase Backend',
        '7. Paystack Payment Integration',
        '8. Email Setup (Zoho Mail)',
        '9. Google Search Console & SEO',
        '10. App Features Reference',
        '11. Premium Tier & Monetization',
        '12. Activation Codes',
        '13. Environment Variables & Secrets',
        '14. Deployment Commands',
        '15. Troubleshooting Guide',
        '16. Future Roadmap',
    ]
    for item in toc_items:
        numbered(doc, item)

    # ── 1. Project Overview ──────────────────────────────────────────────
    heading1(doc, '1. Project Overview')
    normal(doc, 'Oja POS is a point-of-sale system built for Nigerian retail shops and supermarkets. '
                '"Oja" means "market" in Yoruba. The app is designed to work offline-first with optional '
                'cloud sync, supporting Nigerian business workflows including Naira currency, WhatsApp '
                'receipts, and local payment methods.')

    add_table(doc, ['Property', 'Value'], [
        ['App Name', 'Oja POS'],
        ['Tagline', 'The POS Built for Nigerian Shops'],
        ['Landing Page', 'https://ojapos.app'],
        ['Web App', 'https://app.ojapos.app'],
        ['GitHub', 'github.com/Vanka07/oja-pos'],
        ['Brand Color', '#E05E1B (Burnt Orange)'],
        ['Font', 'Poppins'],
        ['Target Market', 'Nigerian retail shops — Lagos first (Alaba, Computer Village, Balogun)'],
        ['App Icon', 'Shopping cart (38 px in-app), bigger and bolder. Tighter gap between cart icon and "Oja" text.'],
        ['Favicon', 'Clean "O" lettermark in orange circle'],
    ], col_widths=[2.0, 4.5])

    # ── 2. Technology Stack ──────────────────────────────────────────────
    heading1(doc, '2. Technology Stack')

    heading2(doc, 'Frontend (Mobile + Web)')
    bullet(doc, 'Expo SDK 53 (React Native)')
    bullet(doc, 'TypeScript')
    bullet(doc, 'NativeWind (Tailwind CSS for React Native)')
    bullet(doc, 'Zustand (state management — 7 stores)')
    bullet(doc, 'React Native Reanimated (animations)')
    bullet(doc, 'MMKV (native storage) / localStorage (web)')
    bullet(doc, 'Expo Router (file-based routing)')

    heading2(doc, 'Backend')
    bullet(doc, 'Supabase (PostgreSQL, Auth, Edge Functions, RLS)')
    bullet(doc, 'Supabase project ID: bjpqdfcpclmcxyydtcmm')
    bullet(doc, 'Region: London (eu-west-2)')

    heading2(doc, 'Hosting & Deployment')
    bullet(doc, 'Vercel (landing page + web app)')
    bullet(doc, 'Vercel project: jamius-projects-ae6688b3/dist')
    bullet(doc, 'Custom domains via Vercel DNS')

    heading2(doc, 'Payments')
    bullet(doc, 'Paystack (Nigerian payment gateway)')
    bullet(doc, 'Supabase Edge Functions for verification + webhooks')

    # ── 3. Repository & Codebase ─────────────────────────────────────────
    heading1(doc, '3. Repository & Codebase')

    heading2(doc, 'Key Directories')
    add_table(doc, ['Directory', 'Description'], [
        ['src/app/', 'All screens (Expo Router file-based routing)'],
        ['src/store/', '7 Zustand stores (retail, auth, staff, subscription, theme, onboarding, update)'],
        ['src/lib/', 'Utilities (storage, sync, paystack, activation, printer, alerts)'],
        ['src/components/', 'Shared components (PremiumUpsell, OjaLogo, TabBar, etc.)'],
        ['supabase/functions/', 'Edge functions (paystack-verify, paystack-webhook)'],
        ['supabase/migrations/', 'SQL migrations'],
        ['scripts/', 'CLI tools (generate-codes.js)'],
        ['dist/', 'Expo web export output (deployed to Vercel)'],
        ['assets/', 'App icon, splash screen, adaptive icon'],
        ['patches/', 'Patch-package fixes'],
        ['oja-landing/ (sibling)', 'Landing page HTML/CSS (separate Vercel project)'],
    ], col_widths=[2.5, 4.0])

    heading2(doc, 'Key Files')
    bullet(doc, 'src/lib/storage.ts — Platform-specific storage adapter (MMKV native, localStorage web)')
    bullet(doc, 'src/lib/syncService.ts — Cloud sync logic (Supabase ↔ local, 5-min interval)')
    bullet(doc, 'src/lib/paystack.ts — Paystack client (initializePayment, verifyPayment)')
    bullet(doc, 'src/lib/activationCode.ts — HMAC-SHA256 code generation + validation')
    bullet(doc, 'src/lib/premiumFeatures.ts — Feature access map + canAccess() helper')
    bullet(doc, 'src/lib/lowStockAlerts.ts — WhatsApp inventory alert system')
    bullet(doc, 'src/lib/printerService.ts — Bluetooth thermal printer + web print fallback')
    bullet(doc, 'src/store/subscriptionStore.ts — Tier management (starter/business)')

    # ── 4. Domain & DNS Setup ────────────────────────────────────────────
    heading1(doc, '4. Domain & DNS Setup')

    heading2(doc, 'Domains')
    add_table(doc, ['Domain', 'Registrar', 'Points To', 'Status'], [
        ['ojapos.app', 'Namecheap (~$13/yr)', 'Landing page (Vercel)', '✅ Live'],
        ['app.ojapos.app', 'Subdomain', 'POS web app (Vercel)', '✅ Live'],
        ['ojapos.ng', 'DomainKing (premium .ng)', 'Landing page (Vercel)', '✅ Live'],
    ])

    heading2(doc, 'DNS Configuration')
    normal(doc, 'Both ojapos.app and ojapos.ng use Vercel nameservers:')
    bullet(doc, 'ns1.vercel-dns.com')
    bullet(doc, 'ns2.vercel-dns.com')
    normal(doc, 'ojapos.app was originally on Namecheap BasicDNS but was switched to Vercel nameservers '
                'to support Zoho Mail MX records + Vercel hosting simultaneously.')

    heading2(doc, 'DNS Records (ojapos.app)')
    add_table(doc, ['Type', 'Name', 'Value', 'Purpose'], [
        ['CNAME', 'app', 'cname.vercel-dns.com', 'POS web app'],
        ['CNAME', 'www', 'cname.vercel-dns.com', 'WWW redirect'],
        ['MX', '@', 'mx.zoho.com (10)', 'Email'],
        ['MX', '@', 'mx2.zoho.com (20)', 'Email fallback'],
        ['MX', '@', 'mx3.zoho.com (50)', 'Email fallback'],
        ['TXT', '@', 'v=spf1 include:zoho.com ~all', 'SPF'],
        ['TXT', '@', 'zoho-verification=...', 'Domain verification'],
        ['CNAME', 'zb...._domainkey', '...zoho.com', 'DKIM'],
    ])

    # ── 5. Vercel Deployment ─────────────────────────────────────────────
    heading1(doc, '5. Vercel Deployment')

    heading2(doc, 'Projects')
    add_table(doc, ['Project', 'Domain', 'Source'], [
        ['dist (jamius-projects-ae6688b3)', 'app.ojapos.app', 'Expo web export'],
        ['oja-landing', 'ojapos.app, ojapos.ng, www.ojapos.app', 'Static HTML'],
    ])

    heading2(doc, 'Deploy POS App')
    normal(doc, 'Run these commands from the oja-pos directory:')
    code_block(doc,
        '# 1. Export web build\n'
        'npx expo export --platform web\n\n'
        '# 2. Copy Vercel config (Expo wipes dist/ each time)\n'
        'cp vercel.json dist/\n\n'
        '# 3. Deploy to production\n'
        'cd dist\n'
        'npx vercel --prod --yes --token "$VERCEL_TOKEN"')

    heading2(doc, 'Deploy Landing Page')
    normal(doc, 'The landing page is in the oja-landing/ directory (sibling to oja-pos):')
    code_block(doc, 'cd oja-landing\nnpx vercel --prod --yes --token "$VERCEL_TOKEN"')

    doc.add_paragraph()
    normal(doc, 'Important: The Vercel token is stored in oja-pos/.env as VERCEL_TOKEN. '
                'Source it before running deploy commands.')

    # ── 6. Supabase Backend ──────────────────────────────────────────────
    heading1(doc, '6. Supabase Backend')

    add_table(doc, ['Property', 'Value'], [
        ['Project ID', 'bjpqdfcpclmcxyydtcmm'],
        ['Region', 'London (eu-west-2)'],
        ['Dashboard', 'https://supabase.com/dashboard/project/bjpqdfcpclmcxyydtcmm'],
        ['API URL', 'https://bjpqdfcpclmcxyydtcmm.supabase.co'],
        ['Auth', 'Email/password (no social login yet)'],
    ], col_widths=[2.0, 4.5])

    heading2(doc, 'Database Tables (8 total)')
    add_table(doc, ['Table', 'Description'], [
        ['shops', 'Shop profiles'],
        ['shop_members', 'Staff linked to shops'],
        ['products', 'Product catalog'],
        ['sales', 'Transaction records'],
        ['customers', 'Customer directory'],
        ['expenses', 'Expense tracking'],
        ['subscriptions', 'Premium tier subscriptions (Paystack + activation codes)'],
        ['sync_log', 'Cloud sync tracking'],
    ], col_widths=[2.5, 4.0])
    normal(doc, 'All tables have Row Level Security (RLS) enabled. Users can only access data belonging to their shop.')

    heading2(doc, 'Edge Functions')
    add_table(doc, ['Function', 'URL', 'JWT'], [
        ['paystack-verify', 'https://bjpqdfcpclmcxyydtcmm.supabase.co/functions/v1/paystack-verify', 'OFF'],
        ['paystack-webhook', 'https://bjpqdfcpclmcxyydtcmm.supabase.co/functions/v1/paystack-webhook', 'OFF'],
    ])

    heading2(doc, 'Supabase Secrets')
    bullet(doc, 'PAYSTACK_SECRET_KEY — Paystack secret key (currently test key)')
    bullet(doc, 'SUPABASE_SERVICE_ROLE_KEY — Service role key for admin operations')

    # ── 7. Paystack Payment Integration ──────────────────────────────────
    heading1(doc, '7. Paystack Payment Integration')

    heading2(doc, 'Overview')
    normal(doc, 'Paystack handles subscription payments for the ₦5,000/month Business tier. The flow is:')
    normal(doc, '1. User taps "Pay with Paystack" in the app')
    normal(doc, '2. App opens Paystack checkout (WebView on mobile, popup on web)')
    normal(doc, '3. User pays via card, bank transfer, or USSD')
    normal(doc, '4. Paystack redirects to callback URL (ojapos.app/payment-callback.html)')
    normal(doc, '5. App calls paystack-verify edge function to confirm payment')
    normal(doc, '6. Edge function verifies with Paystack API + creates subscription record in Supabase')
    normal(doc, '7. Paystack also sends webhook for server-side confirmation (backup)')

    # NEW: Redundancy note
    p = doc.add_paragraph()
    r = p.add_run('Note: ')
    r.bold = True
    r2 = p.add_run('The verify edge function also creates/updates the subscription in Supabase '
                    '(not just the webhook). Both paths write to the subscriptions table for redundancy.')

    heading2(doc, 'Keys (Currently TEST)')
    add_table(doc, ['Key', 'Value'], [
        ['Public Key', 'pk_test_25d612955358dd9ac5ee6f429181c83d2af86816'],
        ['Secret Key', 'Stored in Supabase secrets (PAYSTACK_SECRET_KEY)'],
        ['Webhook URL', 'https://bjpqdfcpclmcxyydtcmm.supabase.co/functions/v1/paystack-webhook'],
    ], col_widths=[2.0, 4.5])

    heading2(doc, 'Going Live Checklist')
    bullet(doc, 'Log into Paystack dashboard (dashboard.paystack.com)')
    bullet(doc, 'Set webhook URL: https://bjpqdfcpclmcxyydtcmm.supabase.co/functions/v1/paystack-webhook')
    # NEW: Mark webhook as done
    p = doc.add_paragraph('Webhook URL has been set in Paystack dashboard ✅', style='List Bullet')
    for r in p.runs:
        r.bold = True
    bullet(doc, 'Switch to Live mode and get live keys')
    bullet(doc, 'Update public key in src/lib/paystack.ts')
    bullet(doc, 'Update secret key in Supabase edge function secrets')
    bullet(doc, 'Deploy payment-callback.html to ojapos.app')
    bullet(doc, 'Test a real ₦100 payment end-to-end')

    # NEW: End-to-end test result
    p = doc.add_paragraph()
    r = p.add_run('✅ End-to-end test completed successfully on Feb 1 2026 — payment processed, '
                   'verify edge function confirmed, subscription record created in Supabase.')
    r.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x7A, 0x22)  # dark green

    # ── 8. Email Setup ───────────────────────────────────────────────────
    heading1(doc, '8. Email Setup (Zoho Mail)')

    add_table(doc, ['Property', 'Value'], [
        ['Email', 'hello@ojapos.app'],
        ['Provider', 'Zoho Mail (free tier)'],
        ['Login', 'https://mail.zoho.com'],
        ['DNS', 'MX + SPF + DKIM configured on Vercel DNS'],
    ], col_widths=[2.0, 4.5])

    normal(doc, 'Used for customer support and business communications. All DNS records '
                '(MX, SPF, DKIM, verification TXT) are set in Vercel DNS for ojapos.app.')

    # ── 9. Google Search Console & SEO ───────────────────────────────────
    heading1(doc, '9. Google Search Console & SEO')

    heading2(doc, 'Search Console')
    bullet(doc, 'ojapos.app — Verified ✅, sitemap submitted')
    bullet(doc, 'ojapos.ng — Verified ✅, sitemap submitted')
    bullet(doc, 'Dashboard: https://search.google.com/search-console')

    heading2(doc, 'SEO Pages')
    add_table(doc, ['URL', 'Description'], [
        ['/', 'Main landing page — hero, features, pricing, testimonials'],
        ['/features', '12 features detailed with descriptions'],
        ['/pricing', 'Starter vs Business tier comparison'],
        ['/faq', '16 questions with FAQPage JSON-LD schema'],
        ['/about', 'Company story, Nigerian market context'],
        ['/privacy', 'Privacy policy'],
    ], col_widths=[1.5, 5.0])

    heading2(doc, 'SEO Features')
    bullet(doc, 'JSON-LD structured data (SoftwareApplication + WebApplication schemas)')
    bullet(doc, 'FAQPage schema for rich results in Google')
    bullet(doc, 'Canonical URLs on all pages')
    bullet(doc, 'sitemap.xml + robots.txt')
    bullet(doc, 'OG meta tags + social preview image')
    bullet(doc, 'Target keywords: "POS app Nigeria", "point of sale Lagos", "retail POS Nigerian shops"')

    # ── 10. App Features Reference ───────────────────────────────────────
    heading1(doc, '10. App Features Reference')

    heading2(doc, 'Core Features (Free Tier)')
    bullet(doc, 'Quick Sell — Scan barcode or tap to add products, cash/transfer/POS payment')
    bullet(doc, 'Product Management — Add/edit products, categories, track stock levels (50 product limit)')
    bullet(doc, 'Sales History — Today\'s transactions with basic totals')
    bullet(doc, 'Credit Book — Track customer debts and payments')
    bullet(doc, 'Customers — Directory with purchase history')
    bullet(doc, 'WhatsApp Receipts — Share receipt via WhatsApp after sale')
    bullet(doc, 'Expenses — Track business expenses by category')
    bullet(doc, 'PIN Lock — Secure app access')

    heading2(doc, 'Business Tier Features (₦5,000/month)')
    bullet(doc, 'Unlimited Products — No 50 product cap')
    bullet(doc, 'Multi-Staff — 4 roles: Owner, Manager, Cashier, Employee')
    bullet(doc, 'Staff PIN Auth — Each staff has unique PIN, activity logged')
    bullet(doc, 'Cloud Sync — Supabase auto-sync every 5 minutes')
    bullet(doc, 'Advanced Reports — Revenue trends, profit margin %, best day, date filtering, charts')
    bullet(doc, 'Payroll — Staff salary tracking, payment recording (Cash/Transfer), status badges')
    bullet(doc, 'WhatsApp Inventory Alerts — Auto-detect low stock after sales, send WhatsApp alert')
    bullet(doc, 'Receipt Printing — Bluetooth thermal printer (58/80mm) + web browser print')
    bullet(doc, 'Export Data — JSON backup via share sheet')
    bullet(doc, 'Dark/Light Mode — System, dark, or light theme toggle')
    bullet(doc, 'Shop Profile — Editable shop name, address, phone, logo')

    heading2(doc, 'UI / Mobile Web')
    # NEW: Tab bar safe area
    bullet(doc, 'Tab bar has mobile web safe area handling — viewport-fit=cover in meta tag, '
                'extra bottom padding to clear the Safari toolbar on iOS. Prevents tab bar buttons '
                'from being hidden behind the browser chrome.')

    heading2(doc, 'Staff Roles & Permissions')
    add_table(doc, ['Role', 'Sell', 'Manage Products', 'View Reports', 'Manage Staff', 'Manage Shop', 'Payroll'], [
        ['Owner', '✅', '✅', '✅', '✅', '✅', '✅'],
        ['Manager', '✅', '✅', '✅', '❌', '❌', '✅'],
        ['Cashier', '✅', '❌', '❌', '❌', '❌', '❌'],
        ['Employee', '—', '—', '—', '—', '—', '— (payroll only, no app access)'],
    ])

    # ── 11. Premium Tier & Monetization ──────────────────────────────────
    heading1(doc, '11. Premium Tier & Monetization')

    add_table(doc, ['Property', 'Value'], [
        ['Free Tier', 'Starter — 50 products, 1 staff, basic reports, offline-only'],
        ['Paid Tier', 'Business — ₦5,000/month, unlimited everything + cloud sync'],
        ['Activation', 'Paystack payment OR manual activation code (OJA-XXXX-XXXX)'],
        ['Gate Status', 'canAccess() currently returns true always (testing bypass)'],
    ], col_widths=[2.0, 4.5])

    normal(doc, 'Note: Premium gates are bypassed for testing. To re-enable, update canAccess() '
                'in src/lib/premiumFeatures.ts to check actual subscription status instead of returning true.')

    # ── 12. Activation Codes ─────────────────────────────────────────────
    heading1(doc, '12. Activation Codes')

    heading2(doc, 'Format')
    normal(doc, 'OJA-XXXX-XXXX (e.g., OJA-A3F7-K9M2)')
    normal(doc, 'Codes are HMAC-SHA256 signed for offline validation — no server call needed to verify authenticity.')

    heading2(doc, 'Generating Codes')
    code_block(doc,
        'cd oja-pos\n'
        'node scripts/generate-codes.js --count 10 --duration 30\n\n'
        '# Options:\n'
        '#   --count N     Number of codes to generate (default: 10)\n'
        '#   --duration N  Subscription duration in days (default: 30)\n'
        '#   --output FILE Output file path')
    doc.add_paragraph()
    normal(doc, 'Generated codes are saved to scripts/codes-YYYY-MM-DD.txt')

    heading2(doc, 'How It Works')
    bullet(doc, 'Code contains encoded duration + HMAC signature')
    bullet(doc, 'App validates signature locally using shared secret')
    bullet(doc, 'Used codes are tracked in local storage to prevent reuse')
    bullet(doc, 'No internet required for validation')

    # ── 13. Environment Variables & Secrets ──────────────────────────────
    heading1(doc, '13. Environment Variables & Secrets')

    heading2(doc, 'Local (.env in oja-pos/)')
    add_table(doc, ['Variable', 'Description'], [
        ['VERCEL_TOKEN', 'Vercel deploy token (full scope)'],
    ], col_widths=[3.0, 3.5])

    heading2(doc, 'Supabase Edge Function Secrets')
    add_table(doc, ['Secret', 'Description'], [
        ['PAYSTACK_SECRET_KEY', 'Paystack secret key (currently test key sk_test_...)'],
        ['SUPABASE_SERVICE_ROLE_KEY', 'Supabase service role key for admin DB operations'],
    ], col_widths=[3.0, 3.5])

    heading2(doc, 'In-App Config (hardcoded)')
    bullet(doc, 'Supabase URL + anon key: src/lib/supabase.ts')
    bullet(doc, 'Paystack public key: src/lib/paystack.ts')
    bullet(doc, 'Activation code secret: src/lib/activationCode.ts')

    # ── 14. Deployment Commands ──────────────────────────────────────────
    heading1(doc, '14. Deployment Commands')

    heading2(doc, 'Deploy POS Web App')
    code_block(doc,
        'cd /path/to/oja-pos\n'
        'source .env\n'
        'npx expo export --platform web\n'
        'cp vercel.json dist/\n'
        'cd dist\n'
        'npx vercel --prod --yes --token "$VERCEL_TOKEN"')

    heading2(doc, 'Deploy Landing Page')
    code_block(doc,
        'cd /path/to/oja-landing\n'
        'npx vercel --prod --yes --token "$VERCEL_TOKEN"')

    heading2(doc, 'Push to GitHub')
    code_block(doc,
        'cd /path/to/oja-pos\n'
        'git add -A\n'
        'git commit -m "your message"\n'
        'git push')

    # ── 15. Troubleshooting Guide ────────────────────────────────────────
    heading1(doc, '15. Troubleshooting Guide')

    heading2(doc, 'Expo web export shows blank page')
    normal(doc, 'Check vercel.json is in dist/ folder. The SPA rewrite rule is required:\n'
                '{"rewrites": [{"source": "/(.*)", "destination": "/index.html"}]}\n'
                'Expo wipes dist/ on every export — always re-copy vercel.json.')

    heading2(doc, 'TextInput only accepts first character (React Native Web)')
    normal(doc, 'NEVER define components inside render functions. This causes React to unmount/remount '
                'the component on every state change. Extract to a separate component file or use a '
                'render function (not a component).')

    heading2(doc, 'WhatsApp link opens blank page on web')
    normal(doc, 'Use window.open("https://wa.me/...") on web, NOT Linking.openURL("whatsapp://..."). '
                'The whatsapp:// deep link protocol doesn\'t work in browsers.')

    heading2(doc, 'Dark mode styles not applying')
    normal(doc, 'Check for duplicate dark: prefixes (e.g., dark:bg-stone-200 dark:bg-stone-800). '
                'NativeWind only processes one dark: prefix per property. Also ensure setColorScheme() '
                'is called in root _layout.tsx.')

    heading2(doc, 'Cloud sync not working')
    normal(doc, 'Check: 1) User is logged in (authStore), 2) Supabase URL/key correct in '
                'src/lib/supabase.ts, 3) RLS policies allow the user\'s shop, 4) Network connectivity.')

    heading2(doc, 'Paystack payment not verifying')
    normal(doc, 'Check: 1) paystack-verify edge function is deployed with JWT OFF, 2) PAYSTACK_SECRET_KEY '
                'is set in Supabase secrets, 3) Payment reference matches, 4) Test vs Live key mismatch.')

    heading2(doc, 'Activation code rejected')
    normal(doc, 'Code may have been used already (tracked in local storage). Check activationCode.ts — '
                'used codes are stored in a Set. To reset for testing, clear the subscription store.')

    heading2(doc, 'Vercel deploy fails')
    normal(doc, 'Ensure VERCEL_TOKEN is valid (check oja-pos/.env). Token must have full scope. '
                'If expired, generate a new one at vercel.com/account/tokens.')

    heading2(doc, 'Staff PIN not working after lock')
    normal(doc, 'The lock screen delegates to staffStore.switchStaff(). Ensure staff PINs are set '
                'in the staff management screen and authStore.authenticate() calls staffStore correctly.')

    heading2(doc, 'Receipt printer not connecting')
    normal(doc, 'Bluetooth thermal printers require native builds (not Expo Go). On web, it falls back '
                'to browser print dialog. Check printer is paired in device Bluetooth settings.')

    # NEW: Subscription not created after payment
    heading2(doc, 'Subscription not created after payment')
    normal(doc, 'Check that SUPABASE_SERVICE_ROLE_KEY is set in Supabase edge function secrets. '
                'Supabase auto-injects it but verify in Dashboard → Edge Functions → Secrets. '
                'The verify function checks if(SUPABASE_SERVICE_ROLE_KEY) before writing — '
                'if empty, it silently skips the DB write.')

    # ── 16. Future Roadmap ───────────────────────────────────────────────
    heading1(doc, '16. Future Roadmap')

    heading2(doc, 'Near Term')
    bullet(doc, 'Set Paystack webhook URL in dashboard + test full payment flow')
    bullet(doc, 'Deploy payment-callback.html to landing site')
    bullet(doc, 'Re-enable premium feature gates (canAccess())')
    bullet(doc, 'Update app icon + splash screen with cart logo + burnt orange')
    bullet(doc, 'Beta test with 5-10 real shop owners in Lagos')

    heading2(doc, 'Medium Term')
    bullet(doc, 'WhatsApp Business API for automated receipts')
    bullet(doc, 'Multi-branch support')
    bullet(doc, 'Supplier Credit Book (BNPL for shops)')
    bullet(doc, 'WhatsApp Storefront integration')
    bullet(doc, 'Agent/Reseller dashboard for distributors')
    bullet(doc, 'Multi-language (Yoruba, Igbo, Hausa, Pidgin)')

    heading2(doc, 'Long Term')
    bullet(doc, 'USSD/SMS fallback for feature phones')
    bullet(doc, 'Price Intelligence (PriceNija integration)')
    bullet(doc, 'Loyalty/Rewards program')
    bullet(doc, 'Supplier ordering system')
    bullet(doc, 'iOS App Store + Google Play Store releases')

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f'✅ Generated: {OUTPUT}')


if __name__ == '__main__':
    build()
